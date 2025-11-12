import os
import sys
import json
import customtkinter as ctk
import sqlite3
from cryptography.fernet import Fernet
from cryptography.hazmat.primitives.kdf.pbkdf2 import PBKDF2HMAC
from cryptography.hazmat.backends import default_backend
from cryptography.hazmat.primitives import hashes
from base64 import urlsafe_b64encode, urlsafe_b64decode
from tkinter import messagebox, StringVar, filedialog
from docx import Document
from odf.opendocument import OpenDocumentText
from odf.text import P

# --- App Icon ---
APP_ICON_PATH = r"Icons\BlackHole_Icon.ico"  # <-- change to your .ico path

# --- Paths ---
local_appdata = os.getenv("LOCALAPPDATA") or os.getenv("APPDATA")
nova_folder = os.path.join(local_appdata, "NovaFoundry")
os.makedirs(nova_folder, exist_ok=True)

db_path = os.path.join(nova_folder, "BlackHolePasswords.db")
settings_path = os.path.join(nova_folder, "settings.json")

# --- Theme (Deep Space Glow) colors ---
BG = "#05050a"
CARD = "#0b0b0f"
CARD_HOVER = "#111327"
ACCENT = "#47a3ff"
ACCENT_DIM = "#2b6f9f"
TEXT = "#e6eef8"

# --- Helper: Derive Key from master password ---
def derive_key(password: str, salt: bytes) -> bytes:
    kdf = PBKDF2HMAC(
        algorithm=hashes.SHA256(),
        length=32,
        salt=salt,
        iterations=390000,
        backend=default_backend()
    )
    return urlsafe_b64encode(kdf.derive(password.encode()))

# --- Initialize DB ---
def init_db():
    conn = sqlite3.connect(db_path)
    c = conn.cursor()
    c.execute('''
        CREATE TABLE IF NOT EXISTS passwords (
            id INTEGER PRIMARY KEY,
            title TEXT,
            username TEXT,
            password TEXT,
            notes TEXT
        )
    ''')
    conn.commit()
    return conn, c

conn, c = init_db()

# --- Password Manager App ---
class PasswordManager(ctk.CTk):
    def __init__(self):
        super().__init__()

        ctk.set_appearance_mode("dark")
        try:
            ctk.set_default_color_theme("dark-blue")
        except Exception:
            pass

        # Hide main window until auth succeeds
        self.title("Black Hole Password Manager")
        self.geometry("900x560")
        if os.path.exists(APP_ICON_PATH):
            try:
                self.iconbitmap(APP_ICON_PATH)
            except Exception:
                pass
        self.attributes('-alpha', 0.0)

        # encryption / state
        self.master_password = None
        self.fernet = None
        self.salt = None

        # load settings
        self.settings = {"master_password_set": False}
        if os.path.exists(settings_path):
            try:
                with open(settings_path, "r") as f:
                    self.settings = json.load(f)
                if "salt" in self.settings:
                    self.salt = urlsafe_b64decode(self.settings["salt"])
            except Exception:
                self.settings = {"master_password_set": False}

        # show master password popup
        success = self._show_master_password_modal()
        if not success:
            self.destroy()
            sys.exit()

        # Build UI after successful auth
        self.configure(fg_color=BG)
        self._build_ui()
        self.attributes('-alpha', 1.0)
        self.load_cards()

    # --- Auth modal ---
    def _show_master_password_modal(self):
        popup = ctk.CTkToplevel(self)
        popup.grab_set()
        popup.configure(fg_color=BG)
        popup.resizable(False, False)
        popup.attributes("-topmost", True)

        # Center popup
        width, height = 480, 320
        x = (popup.winfo_screenwidth() // 2) - (width // 2)
        y = (popup.winfo_screenheight() // 2) - (height // 2)
        popup.geometry(f"{width}x{height}+{x}+{y}")

        closed_by_user = {"val": True}

        def on_close():
            popup.grab_release()
            popup.destroy()

        popup.protocol("WM_DELETE_WINDOW", on_close)

        # Header
        ctk.CTkLabel(popup, text="Black Hole — Master Password",
                     font=("Helvetica", 16, "bold"),
                     text_color=TEXT, fg_color=BG).pack(pady=(16,6))
        ctk.CTkLabel(popup, text="Protect your vault with a master password",
                     text_color=ACCENT_DIM, fg_color=BG).pack(pady=(0,12))

        # Entry frame
        frame = ctk.CTkFrame(popup, fg_color=CARD, corner_radius=8)
        frame.pack(padx=20, pady=8, fill="both", expand=False)

        pwd_var = StringVar()
        pwd_entry = ctk.CTkEntry(frame, placeholder_text="Master Password",
                                 show="*", textvariable=pwd_var,
                                 width=360, justify="center")
        pwd_entry.pack(padx=12, pady=(12,6))
        
        # --- ADDED THIS LINE ---
        original_entry_color = pwd_entry.cget("fg_color") 

        def toggle_pwd():
            pwd_entry.configure(show="" if pwd_entry.cget("show")=="*" else "*")

        ctk.CTkButton(frame, text="Show/Hide", width=120,
                       command=toggle_pwd, fg_color=ACCENT, text_color=BG,
                       hover_color=ACCENT_DIM).pack(pady=(0,10))

        # Optional save path
        path_var = StringVar()
        def browse_path():
            file_path = filedialog.asksaveasfilename(defaultextension=".txt",
                                                     filetypes=[("Text files","*.txt")])
            if file_path:
                path_var.set(file_path)
        ctk.CTkButton(frame, text="Select Save Path (optional)", command=browse_path,
                       fg_color=ACCENT, text_color=BG, hover_color=ACCENT_DIM).pack(pady=(0,8))
        ctk.CTkLabel(frame, textvariable=path_var, text_color=TEXT, fg_color=CARD).pack(pady=(0,8))

        # Buttons
        btn_frame = ctk.CTkFrame(popup, fg_color=BG, corner_radius=0)
        btn_frame.pack(pady=(12,12))

        # --- REPLACED THIS FUNCTION ---
        def mark_wrong(): 
            # 1. Turn entry red
            pwd_entry.configure(fg_color="#7a2d2d")

            # 2. Get original position
            try:
                # Geometry string is 'widthxheight+x+y'
                geo_string = popup.geometry()
                parts = geo_string.split('+')
                size_part = parts[0] # 'widthxheight'
                original_x = int(parts[1])
                original_y = int(parts[2])
            except Exception:
                # Failsafe if geometry string is weird or window closed
                pwd_entry.configure(fg_color="#7a2d2d")
                return 

            # 3. Define the shake animation
            def shake_animation(step=0):
                try:
                    # Movements are offsets from the original_x
                    offsets = [10, -10, 10, -10, 5, -5, 0] 
                    
                    if step < len(offsets):
                        dx = offsets[step]
                        # Apply new position
                        popup.geometry(f"{size_part}+{original_x + dx}+{original_y}")
                        # Schedule next step
                        popup.after(50, shake_animation, step + 1)
                    
                    elif step == len(offsets):
                        # After shake is done, wait 1 sec and reset color
                        popup.after(1000, lambda: pwd_entry.configure(fg_color=original_entry_color))
                
                except Exception:
                    # Failsafe in case popup was destroyed during animation
                    pass

            # Start the shake
            shake_animation(0)

        def create_master():
            pwd = pwd_var.get() or ""
            if not pwd:
                mark_wrong()
                messagebox.showerror("Error", "Master password required!", parent=popup)
                return
            try:
                self.salt = os.urandom(16)
                key = derive_key(pwd, self.salt)
                self.fernet = Fernet(key)
                verif_enc = self.fernet.encrypt(b"VERIFICATION").decode()
                self.settings["salt"] = urlsafe_b64encode(self.salt).decode()
                self.settings["verification"] = verif_enc
                self.master_password = pwd
            except Exception as e:
                mark_wrong()
                messagebox.showerror("Error", f"Failed to derive key: {e}", parent=popup)
                return
            save_path = path_var.get()
            if save_path:
                try:
                    with open(save_path,"w",encoding="utf-8") as f:
                        f.write(pwd)
                    messagebox.showinfo("Saved", f"Master password saved to {save_path}", parent=popup)
                except Exception as e:
                    messagebox.showwarning("Warning", f"Couldn't save: {e}", parent=popup)
            self.settings["master_password_set"] = True
            try:
                with open(settings_path,"w",encoding="utf-8") as sf:
                    json.dump(self.settings,sf)
            except Exception:
                pass
            closed_by_user["val"] = False
            popup.grab_release()
            popup.destroy()

        def unlock_master():
            pwd = pwd_var.get() or ""
            if not pwd:
                mark_wrong()
                messagebox.showerror("Error", "Master password required!", parent=popup)
                return
            try:
                key = derive_key(pwd, self.salt)
                fernet_test = Fernet(key)
                verif_dec = fernet_test.decrypt(self.settings["verification"].encode()).decode()
                if verif_dec != "VERIFICATION":
                    raise ValueError("Verification failed")
                self.fernet = fernet_test
                self.master_password = pwd
                closed_by_user["val"] = False
                popup.grab_release()
                popup.destroy()
            except Exception:
                mark_wrong()
                messagebox.showerror("Error", "Incorrect master password!", parent=popup)

        if not self.settings.get("master_password_set", False):
            ctk.CTkButton(btn_frame, text="Create & Continue", command=create_master,
                           fg_color=ACCENT, text_color=BG, hover_color=ACCENT_DIM, width=200).pack(side="left", padx=12)
        else:
            ctk.CTkButton(btn_frame, text="Unlock", command=unlock_master,
                           fg_color=ACCENT, text_color=BG, hover_color=ACCENT_DIM, width=200).pack(side="left", padx=12)
        ctk.CTkButton(btn_frame, text="Exit", command=on_close, fg_color="#3a3a3a", width=120).pack(side="left", padx=6)

        pwd_entry.focus_set()
        self.wait_window(popup)
        return not closed_by_user["val"] and self.fernet is not None

    # --- Build UI ---
    def _build_ui(self):
        header = ctk.CTkFrame(self, fg_color=BG, height=64)
        header.pack(fill="x", padx=12, pady=(12,0))
        ctk.CTkLabel(header, text="Black Hole Vault", font=("Helvetica", 18, "bold"),
                     text_color=TEXT, fg_color=BG).pack(side="left", padx=(6,12))
        ctk.CTkLabel(header, text="Deep Space Glow", text_color=ACCENT_DIM, fg_color=BG).pack(side="left")

        self.cards_frame = ctk.CTkScrollableFrame(self, width=860, height=420, fg_color=BG, corner_radius=10)
        self.cards_frame.pack(padx=12, pady=12, fill="both", expand=True)

        btn_frame = ctk.CTkFrame(self, fg_color=BG)
        btn_frame.pack(pady=(0,12))
        ctk.CTkButton(btn_frame, text="Create New", command=self.create_new_card,
                       fg_color=ACCENT, text_color=BG, hover_color=ACCENT_DIM).pack(side="left", padx=8)
        ctk.CTkButton(btn_frame, text="Export DOCX", command=self.export_docx, fg_color="#2f3b4a").pack(side="left", padx=8)
        ctk.CTkButton(btn_frame, text="Export ODT", command=self.export_odt, fg_color="#2f3b4a").pack(side="left", padx=8)

    # --- Load Cards ---
    def load_cards(self):
        for widget in self.cards_frame.winfo_children():
            widget.destroy()

        for row in c.execute("SELECT id, title, username, password, notes FROM passwords ORDER BY id DESC"):
            id_, title, user, pwd_enc, notes = row
            try:
                pwd = self.fernet.decrypt(pwd_enc.encode()).decode() if pwd_enc else ""
            except Exception:
                pwd = ""

            card = ctk.CTkFrame(self.cards_frame, fg_color=CARD, corner_radius=10)
            card.pack(pady=8, padx=12, fill="x")

            def on_enter(e, f=card):
                f.configure(fg_color=CARD_HOVER)
            def on_leave(e, f=card):
                f.configure(fg_color=CARD)
            card.bind("<Enter>", on_enter)
            card.bind("<Leave>", on_leave)

            left = ctk.CTkFrame(card, fg_color=CARD, corner_radius=0)
            left.pack(side="left", fill="both", expand=True, padx=8, pady=8)
            ctk.CTkLabel(left, text=title or "(No title)", anchor="w",
                         font=("Helvetica", 14, "bold"), text_color=TEXT, fg_color=CARD).pack(anchor="w")
            ctk.CTkLabel(left, text=f"Username: {user or ''}", anchor="w",
                         text_color=ACCENT_DIM, fg_color=CARD).pack(anchor="w", pady=(4,0))

            pwd_var = StringVar(value="*"*len(pwd) if pwd else "")
            ctk.CTkLabel(left, textvariable=pwd_var, anchor="w", text_color=TEXT, fg_color=CARD).pack(anchor="w", pady=(8,0))

            right = ctk.CTkFrame(card, fg_color=CARD, corner_radius=0)
            right.pack(side="right", padx=10, pady=8)

            def toggle_show(pw=pwd, var=pwd_var):
                var.set(pw if var.get().startswith("*") else "*"*len(pw))
            ctk.CTkButton(right, text="Show", command=toggle_show,
                           width=72, fg_color=ACCENT, text_color=BG, hover_color=ACCENT_DIM).pack(pady=4)
            ctk.CTkButton(right, text="Edit", command=lambda id=id_: self.edit_card_popup(id), width=72).pack(pady=4)
            ctk.CTkButton(right, text="Delete", command=lambda id=id_: self.delete_card(id), width=72,
                           fg_color="#7a2d2d").pack(pady=4)

    # --- Create Card ---
    def create_new_card(self):
        popup = ctk.CTkToplevel(self)
        popup.grab_set()
        popup.geometry("420x220")
        popup.title("Create New Password")
        popup.configure(fg_color=BG)

        ctk.CTkLabel(popup, text="Create New Entry", font=("Helvetica", 14, "bold"), text_color=TEXT, fg_color=BG).pack(pady=(12,6))
        title_var = StringVar()
        ctk.CTkEntry(popup, placeholder_text="Title (required)", textvariable=title_var, width=360).pack(pady=8)

        def create_card_action():
            title = title_var.get().strip()
            if not title:
                messagebox.showerror("Error", "Title required!", parent=popup)
                return
            c.execute("INSERT INTO passwords (title, username, password, notes) VALUES (?, ?, ?, ?)", (title, "", "", ""))
            conn.commit()
            popup.grab_release()
            popup.destroy()
            self.load_cards()

        ctk.CTkButton(popup, text="Create", command=create_card_action,
                       fg_color=ACCENT, text_color=BG, hover_color=ACCENT_DIM, width=140).pack(pady=(12,10))

    # --- Edit Card Popup ---
    def edit_card_popup(self, id_):
        row = c.execute("SELECT title, username, password, notes FROM passwords WHERE id=?", (id_,)).fetchone()
        if not row:
            messagebox.showerror("Error", "Entry not found.")
            return
        title, user, pwd_enc, notes = row
        try:
            pwd = self.fernet.decrypt(pwd_enc.encode()).decode() if pwd_enc else ""
        except Exception:
            pwd = ""

        popup = ctk.CTkToplevel(self)
        popup.grab_set()
        popup.geometry("480x420")
        popup.title("Edit Entry")
        popup.configure(fg_color=BG)

        ctk.CTkLabel(popup, text="Edit Entry", font=("Helvetica", 14, "bold"), text_color=TEXT, fg_color=BG).pack(pady=(12,6))

        ctk.CTkLabel(popup, text="Title", text_color=ACCENT_DIM, fg_color=BG).pack(anchor="w", padx=20)
        title_var = StringVar(value=title)
        ctk.CTkEntry(popup, textvariable=title_var, width=420).pack(padx=20, pady=(4,8))

        ctk.CTkLabel(popup, text="Username", text_color=ACCENT_DIM, fg_color=BG).pack(anchor="w", padx=20)
        user_var = StringVar(value=user)
        ctk.CTkEntry(popup, textvariable=user_var, width=420).pack(padx=20, pady=(4,8))

        ctk.CTkLabel(popup, text="Password", text_color=ACCENT_DIM, fg_color=BG).pack(anchor="w", padx=20)
        pwd_var = StringVar(value=pwd)
        pwd_entry = ctk.CTkEntry(popup, textvariable=pwd_var, show="*", width=420)
        pwd_entry.pack(padx=20, pady=(4,8))
        def toggle_pwd_entry():
            pwd_entry.configure(show="" if pwd_entry.cget("show")=="*" else "*")
        ctk.CTkButton(popup, text="Show/Hide", command=toggle_pwd_entry, fg_color=ACCENT, text_color=BG).pack(pady=(0,8))

        ctk.CTkLabel(popup, text="Notes", text_color=ACCENT_DIM, fg_color=BG).pack(anchor="w", padx=20)
        notes_var = StringVar(value=notes)
        ctk.CTkEntry(popup, textvariable=notes_var, width=420).pack(padx=20, pady=(4,12))

        def save_card():
            enc_pwd = self.fernet.encrypt(pwd_var.get().encode()).decode() if pwd_var.get() else ""
            c.execute("UPDATE passwords SET title=?, username=?, password=?, notes=? WHERE id=?",
                      (title_var.get(), user_var.get(), enc_pwd, notes_var.get(), id_))
            conn.commit()
            popup.grab_release()
            popup.destroy()
            self.load_cards()

        ctk.CTkButton(popup, text="Save", command=save_card, fg_color=ACCENT, text_color=BG, width=120).pack(pady=(6,12))

    # --- Delete ---
    def delete_card(self, id_):
        if messagebox.askyesno("Confirm Delete", "Are you sure you want to delete this entry?"):
            c.execute("DELETE FROM passwords WHERE id=?", (id_,))
            conn.commit()
            self.load_cards()

    # --- Export ---
    def export_docx(self):
        doc = Document()
        for row in c.execute("SELECT title, username, password, notes FROM passwords"):
            title, user, pwd_enc, notes = row
            try:
                pwd = self.fernet.decrypt(pwd_enc.encode()).decode() if pwd_enc else ""
            except Exception:
                pwd = ""
            doc.add_paragraph(f"Title: {title}")
            doc.add_paragraph(f"Username: {user}")
            doc.add_paragraph(f"Password: {pwd}")
            doc.add_paragraph(f"Notes: {notes}")
            doc.add_paragraph("-"*30)
        path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word","*.docx")])
        if path:
            doc.save(path)
            messagebox.showinfo("Exported", f"Exported to {path}")

    def export_odt(self):
        odt = OpenDocumentText()
        for row in c.execute("SELECT title, username, password, notes FROM passwords"):
            title, user, pwd_enc, notes = row
            try:
                pwd = self.fernet.decrypt(pwd_enc.encode()).decode() if pwd_enc else ""
            except Exception:
                pwd = ""
            for text in [f"Title: {title}", f"Username: {user}", f"Password: {pwd}", f"Notes: {notes}", "-"*30]:
                p = P(text=text)
                odt.text.addElement(p)
        path = filedialog.asksaveasfilename(defaultextension=".odt", filetypes=[("OpenDocument","*.odt")])
        if path:
            odt.save(path)
            messagebox.showinfo("Exported", f"Exported to {path}")

# --- Run App ---
if __name__ == "__main__":
    app = PasswordManager()
    app.mainloop()